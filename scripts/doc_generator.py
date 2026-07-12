#!/usr/bin/env python3
"""
环保执法文书生成器
支持生成：行政处罚决定书、责令改正通知书、现场检查笔录、调查询问笔录、
         查封扣押决定书、听证告知书、案件调查报告 等7种执法文书
"""

import os, json, re
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
CASES_FILE = os.path.join(os.path.dirname(SKILL_DIR), 'db', 'cases.json')

# ========== 工具函数 ==========

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1B4332'):
    """添加标题"""
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_para(doc, text, bold=False, indent=0, spacing_after=6):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_table_row(table, data, bold=False, bg=None):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold
                r.font.size = Pt(10.5)
        if bg:
            set_cell_bg(cell, bg)
    return row

def load_cases():
    try:
        with open(CASES_FILE) as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
            return data.get('cases', []) if isinstance(data, dict) else []
    except:
        return []

def search_similar_cases(keyword, case_type=None, limit=3):
    """搜索相似案例"""
    cases = load_cases()
    k = keyword.lower().strip()
    results = []
    for c in cases:
        if case_type and case_type not in (c.get('type','') or ''):
            continue
        # 空关键词返回所有
        if not k:
            results.append((1, c))
            continue
        score = 0
        t = ' '.join([
            c.get('title',''), c.get('fact',''), c.get('type',''),
            c.get('party',''), c.get('result','')
        ]).lower()
        for kw in k.split():
            if kw in t: score += 1
        if score > 0:
            results.append((score, c))
    results.sort(key=lambda x: -x[0])
    return [c for _, c in results[:limit]]

def get_law_text(law_name):
    """获取法规条文文本（简化版，实际应接入知识库）"""
    law_db = {
        '水污染防治法第83条': '《水污染防治法》第八十三条：利用渗井、渗坑、裂隙、溶洞，私设暗管，篡改、伪造监测数据，或者不正常运行水污染防治设施等逃避监管的方式排放水污染物的，由县级以上人民政府环境保护主管部门责令改正或者责令限制生产、停产整治，并处十万元以上一百万元以下的罚款；情节严重的，报经有批准权的人民政府批准，责令停业、关闭。',
        '固体废物法第112条': '《固体废物污染环境防治法》第一百一十二条：违反本法规定，有下列行为之一的，由生态环境主管部门责令改正，处以罚款……（四）擅自倾倒、堆放、丢弃、遗撒工业固体废物，或者未采取相应防范措施，造成工业固体废物扬散、流失、渗漏或者其他环境污染的……处所需处置费用一倍以上三倍以下的罚款，所需处置费用不足三万元的，处三万元以上十万元以下的罚款。',
        '大气污染防治法第99条': '《大气污染防治法》第九十九条：违反本法规定，有下列行为之一的，由县级以上人民政府生态环境主管部门责令改正或者限制生产、停产整治，处十万元以上一百万元以下的罚款；情节严重的，报经有批准权的人民政府批准，责令停业、关闭：（一）未依法取得排污许可证排放大气污染物的；（二）超过大气污染物排放标准或者超过重点大气污染物排放总量控制指标排放大气污染物的；（三）通过逃避监管的方式排放大气污染物的。',
        '刑法第338条': '《刑法》第三百三十八条【污染环境罪】：违反国家规定，排放、倾倒或者处置有放射性的废物、含传染病病原体的废物、有毒物质或者其他有害物质，严重污染环境的，处三年以下有期徒刑或者拘役，并处或者单处罚金；后果特别严重的，处三年以上七年以下有期徒刑，并处罚金。',
        '刑法第286条之一': '《刑法》第二百八十六条之一【破坏计算机信息系统罪】：违反国家规定，对计算机信息系统功能进行删除、修改、增加、干扰，造成计算机信息系统不能正常运行，后果严重的，处五年以下有期徒刑或者拘役；后果特别严重的，处五年以上有期徒刑。',
        '环评法第31条': '《环境影响评价法》第三十一条：建设单位未依法报批建设项目环境影响报告书、报告表，擅自开工建设的，由县级以上生态环境主管部门责令停止建设，根据违法情节和危害后果，处建设项目总投资额百分之一以上百分之五以下的罚款，并可以责令恢复原状。',
        '建设项目环保条例第23条': '《建设项目环境保护管理条例》第二十三条：违反本条例规定，需要配套建设的环境保护设施未建成、未经验收或者验收不合格，建设项目即投入生产或者使用，或者在环境保护设施验收中弄虚作假的，由县级以上环境保护行政主管部门责令限期改正，处20万元以上100万元以下的罚款；逾期不改正的，处100万元以上200万元以下的罚款；对直接负责的主管人员和其他责任人员，处5万元以上20万元以下的罚款。',
        '排污许可条例第33条': '《排污许可管理条例》第三十三条：违反本条例规定排污单位有下列行为之一的，由生态环境主管部门责令改正或者限制生产、停产整治，处20万元以上100万元以下的罚款；情节严重的，报经有批准权的人民政府批准，责令停业、关闭：（一）未取得排污许可证排放污染物的；（二）排污许可证有效期届满未申请延续或者延续申请未经批准排放污染物的。',
        '噪声污染防治法第62条': '《噪声污染防治法》第六十二条：违反本法规定，有下列行为之一的，由地方人民政府指定的部门责令改正，处五千元以上五万元以下的罚款；拒不改正的，由地方人民政府指定的部门依法采取强制措施：（一）超过噪声排放标准排放工业噪声的。',
    }
    for k, v in law_db.items():
        if law_name and law_name in k:
            return v
    return f'法规：{law_name}（请在知识库中查阅完整条文）'

def search_laws(query):
    """搜索相关法规条文"""
    q = query.lower()
    law_db = {
        '水污染防治法': ['第83条（私设暗管/超标排放）', '第84条（逾期不改正）'],
        '大气污染防治法': ['第99条（超标排放/数据造假）', '第100条（不正常运行治污设施）', '第108条（VOCs无组织排放）'],
        '固体废物污染环境防治法': ['第112条（非法倾倒/处置固废）', '第113条（危废台账虚假）'],
        '环境影响评价法': ['第31条（未批先建）', '第32条（越权审批）'],
        '建设项目环境保护管理条例': ['第23条（未验先投）'],
        '排污许可管理条例': ['第33条（无证排污）', '第34条（超许可排放）'],
        '刑法': ['第338条（污染环境罪）', '第286条之一（破坏计算机信息系统罪）'],
        '噪声污染防治法': ['第62条（工业噪声超标）'],
        '土壤污染防治法': ['第87条（土壤污染）'],
        '放射性污染防治法': ['相关条款（辐射类违法）'],
    }
    results = {}
    for law, articles in law_db.items():
        if any(k in law.lower() or k in q for k in q.split()):
            results[law] = articles
    if not results:
        for law, articles in law_db.items():
            if any(a in q or q in law for a in articles):
                results[law] = articles
    return results

# ========== 文书生成函数 ==========

def generate_case_id(doc_type):
    today = date.today()
    count_file = os.path.join(os.path.dirname(SKILL_DIR), 'db', 'case-id-counter.txt')
    try:
        with open(count_file) as f:
            last_date, counter = f.read().strip().split(',')
        counter = int(counter)
    except:
        last_date, counter = '20260101', 0
    if last_date != today.strftime('%Y%m%d'):
        counter = 0
    counter += 1
    with open(count_file, 'w') as f:
        f.write(f'{today.strftime("%Y%m%d")},{counter}')
    return f'{today.strftime("%Y%m%d")}-{counter:04d}'

def create_styled_table(doc, headers, rows, col_widths=None):
    """创建标准表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        set_cell_bg(hdr.cells[i], '1B4332')
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255,255,255)
                r.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = str(text)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def generate_xzcfjds(party, fact, law, result, fine, case_id=None):
    """生成行政处罚决定书"""
    doc = Document()
    cid = case_id or generate_case_id('xzcfjds')
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('行 政 处 罚 决 定 书')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    # 文号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'济环罚字〔{date.today().year}〕第{cid.split("-")[-1]}号')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string('666666')

    doc.add_paragraph()

    # 当事人信息
    p = doc.add_paragraph()
    p.add_run('当事人：').bold = True
    p.add_run(party)
    p.paragraph_format.space_after = Pt(8)

    add_para(doc, f'案号：{cid}')

    # 违法事实
    add_heading(doc, '一、违法事实', level=2)
    add_para(doc, fact)

    # 证据
    add_heading(doc, '二、证据', level=2)
    add_para(doc, '上述违法事实有以下证据予以证实：')
    add_para(doc, '1. 现场检查（勘察）笔录及影像资料', indent=1)
    add_para(doc, '2. 调查询问笔录', indent=1)
    add_para(doc, '3. 监测报告及资质证明', indent=1)
    add_para(doc, '4. 书证材料（营业执照、环评批复等）', indent=1)

    # 违反法律
    add_heading(doc, '三、违反法律', level=2)
    add_para(doc, f'上述行为违反了{law}之规定。')

    # 处罚依据
    add_heading(doc, '四、处罚依据', level=2)
    add_para(doc, f'依据{law}之规定，对你（单位）作出如下行政处罚：')

    # 处罚内容（表格）
    doc.add_paragraph()
    table = create_styled_table(doc, ['处罚项目', '内容'], [
        ['罚款金额', f'人民币 {fine} 万元整（¥{float(fine)*10000:.0f}元）'],
        ['缴纳期限', f'自收到本决定书之日起15日内缴纳至指定账户'],
        ['缴纳方式', '凭本决定书到济南市生态环境局财务室开具缴款通知书'],
    ])
    doc.add_paragraph()

    # 救济途径
    add_heading(doc, '五、救济途径', level=2)
    add_para(doc, f'你（单位）如不服本处罚决定，可在收到本决定书之日起60日内向济南市人民政府申请行政复议，或在6个月内向有管辖权的人民法院提起行政诉讼。')
    add_para(doc, '逾期不申请行政复议、不提起行政诉讼，又不履行本处罚决定的，本机关将依法申请人民法院强制执行。')

    # 署名
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'济南市生态环境局（印章）')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today_str = date.today().strftime('%Y年%m月%d日')
    p.add_run(today_str)

    return doc, cid

def generate_zlgztz(party, fact, requirement, deadline, case_id=None):
    """生成责令改正违法行为决定书"""
    doc = Document()
    cid = case_id or generate_case_id('zlgztz')
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('责令改正违法行为决定书')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'济环责改字〔{date.today().year}〕第{cid.split("-")[-1]}号')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string('666666')

    doc.add_paragraph()
    add_para(doc, f'当事人：{party}')
    add_para(doc, f'案号：{cid}')

    add_heading(doc, '一、违法事实', level=2)
    add_para(doc, fact)

    add_heading(doc, '二、责令改正内容', level=2)
    add_para(doc, requirement)

    add_heading(doc, '三、改正期限', level=2)
    add_para(doc, f'限于 {deadline} 前完成整改。如逾期未改正，我局将依法启动按日连续处罚程序或采取进一步的行政强制措施。')

    add_heading(doc, '四、后续处理', level=2)
    add_para(doc, '整改完成后，你单位应向我局提交整改报告及相关证明材料。我局将依法进行核查。')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('济南市生态环境局（印章）')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(date.today().strftime('%Y年%m月%d日'))

    return doc, cid

def generate_xcjcbcjl(party, location, inspector, date_str, facts, law):
    """生成现场检查笔录"""
    doc = Document()
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('现 场 检 查（勘 察）笔 录')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    fields = [
        ('当事人', party),
        ('检查时间', date_str),
        ('检查地点', location),
        ('执法人员', inspector),
        ('记录人', '（记录人签名）'),
        ('天气状况', '（填写当日天气）'),
    ]
    for i, (k, v) in enumerate(fields):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        set_cell_bg(table.rows[i].cells[0], 'D8E8DC')
        for p_cell in table.rows[i].cells[0].paragraphs:
            for r in p_cell.runs: r.bold = True

    doc.add_paragraph()
    add_heading(doc, '检查情况', level=2)
    add_para(doc, facts)

    add_heading(doc, '现场取证', level=2)
    add_para(doc, '□ 现场照片（  ）张  □ 现场录像（  ）段  □ 采样（  ）个')
    add_para(doc, '□ 书证材料（  ）份  □ 其他：___________________')

    add_heading(doc, '涉嫌违法', level=2)
    add_para(doc, f'上述行为涉嫌违反{law}之规定。')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'被检查单位负责人（签名）：_______________    {date.today().strftime("%Y年%m月%d日")}')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'执法人员（签名）：_______________    {date.today().strftime("%Y年%m月%d日")}')

    return doc, f'XCJC-{date.today().strftime("%Y%m%d")}'

def generate_dcwxblj(party, person, role, questions_answers, case_id=None):
    """生成调查询问笔录"""
    doc = Document()
    cid = case_id or f'DCXW-{date.today().strftime("%Y%m%d")}'
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('调 查 询 问 笔 录')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    fields = [
        ('被询问人', person),
        ('工作单位/职务', '（填写）'),
        ('询问时间', date.today().strftime('%Y年%m月%d日  时 分')),
        ('询问地点', '（填写）'),
    ]
    for i, (k, v) in enumerate(fields):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        set_cell_bg(table.rows[i].cells[0], 'D8E8DC')

    doc.add_paragraph()
    add_heading(doc, '告知事项', level=2)
    add_para(doc, '根据《中华人民共和国行政处罚法》第五十五条，执法人员向被询问人告知：被询问人对询问有如实回答的义务，对与本案无关的问题有拒绝回答的权利；执法人员应如实制作笔录，被询问人有权核对笔录，如实陈述。')

    add_heading(doc, '询问内容', level=2)
    for i, (q, a) in enumerate(questions_answers, 1):
        p = doc.add_paragraph()
        p.add_run(f'问{i}：{q}').bold = True
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        p.add_run(f'答{i}：{a}')
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('被询问人（签名）：_______________    执法人员（签名）：_______________')

    return doc, cid

def generate_cfkyjds(party, fact, legal_basis, measures, case_id=None):
    """生成查封扣押决定书"""
    doc = Document()
    cid = case_id or generate_case_id('cfkyjds')
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('查封（扣押）决定书')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'济环查扣字〔{date.today().year}〕第{cid.split("-")[-1]}号')

    doc.add_paragraph()
    add_para(doc, f'当事人：{party}')
    add_para(doc, f'案号：{cid}')

    add_heading(doc, '一、违法事实', level=2)
    add_para(doc, fact)

    add_heading(doc, '二、查封/扣押依据', level=2)
    add_para(doc, f'依据{legal_basis}之规定，对你（单位）的以下设施、设备、场所实施查封（扣押）：')

    add_heading(doc, '三、查封/扣押清单', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for cell, h in zip(hdr.cells, ['序号', '名称', '规格/型号', '数量']):
        cell.text = h
        set_cell_bg(cell, '1B4332')
        for p_cell in cell.paragraphs:
            for r in p_cell.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255,255,255)
    for i, item in enumerate(measures, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = item
        row.cells[2].text = '—'
        row.cells[3].text = '—'

    add_heading(doc, '四、存放地点', level=2)
    add_para(doc, '查封（扣押）物品存放于：___________________（场所）')

    add_heading(doc, '五、救济途径', level=2)
    add_para(doc, '如不服本决定，可在收到本决定书之日起60日内向济南市人民政府申请行政复议，或在3个月内向有管辖权的人民法院提起行政诉讼。')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('济南市生态环境局（印章）')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(date.today().strftime('%Y年%m月%d日'))

    return doc, cid

def generate_ajdcbg(party, case_type, fact, evidence, law, suggestion, case_id=None):
    """生成案件调查报告"""
    doc = Document()
    cid = case_id or generate_case_id('ajdcbg')
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('案 件 调 查 报 告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1B4332')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'济环案调字〔{date.today().year}〕第{cid.split("-")[-1]}号')

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    fields = [
        ('案件编号', cid),
        ('当事人', party),
        ('案件类型', case_type),
        ('调查日期', date.today().strftime('%Y年%m月%d日')),
        ('承办机构', '济南市生态环境局执法支队'),
    ]
    for i, (k, v) in enumerate(fields):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        set_cell_bg(table.rows[i].cells[0], 'D8E8DC')

    add_heading(doc, '一、案情简介', level=2)
    add_para(doc, fact)

    add_heading(doc, '二、调查经过', level=2)
    add_para(doc, '1. 2026年  月  日，执法人员依法对当事人进行现场检查（勘察）。')
    add_para(doc, '2. 2026年  月  日，对当事人及相关人员进行调查询问。')
    add_para(doc, '3. 2026年  月  日，委托检测机构出具检测报告（报告编号：      ）。')

    add_heading(doc, '三、证据材料', level=2)
    add_para(doc, f'（一）主体证据：{evidence[0]}')
    add_para(doc, f'（二）事实证据：{evidence[1]}')
    add_para(doc, f'（三）鉴定意见：{evidence[2]}')

    add_heading(doc, '四、法律适用', level=2)
    add_para(doc, f'当事人行为违反{law}之规定，应依据{law}予以处罚。')

    add_heading(doc, '五、承办意见', level=2)
    add_para(doc, suggestion)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('调查人员（签名）：___________________')
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('日    期：' + date.today().strftime('%Y年%m月%d日'))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('审核人（签名）：___________________')

    return doc, cid

# ========== 主函数 ==========

def generate_doc(doc_type, case_data, output_path):
    """生成文书主函数"""
    generators = {
        'xzcfjds': lambda d: generate_xzcfjds(d['party'], d['fact'], d['law'], d.get('result',''), d.get('fine', '10')),
        'zlgztz': lambda d: generate_zlgztz(d['party'], d['fact'], d.get('requirement',''), d.get('deadline','')),
        'xcjcbcjl': lambda d: generate_xcjcbcjl(d['party'], d.get('location',''), d.get('inspector',''), d.get('date',''), d.get('facts',''), d.get('law','')),
        'dcwxblj': lambda d: generate_dcwxblj(d['party'], d.get('person',''), d.get('role',''), d.get('qa',[])),
        'cfkyjds': lambda d: generate_cfkyjds(d['party'], d.get('fact',''), d.get('legal_basis',''), d.get('items',[])),
        'ajdcbg': lambda d: generate_ajdcbg(d['party'], d.get('case_type',''), d.get('fact',''), d.get('evidence',[]), d.get('law',''), d.get('suggestion','')),
    }

    if doc_type not in generators:
        raise ValueError(f'未知文书类型：{doc_type}')

    doc, cid = generators[doc_type](case_data)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path, cid

def main():
    import sys
    if len(sys.argv) < 3:
        print('用法: python3 doc_generator.py <文书类型> <输出路径> [JSON参数]')
        print('文书类型: xzcfjds, zlgztz, xcjcbcjl, dcwxblj, cfkyjds, ajdcbg')
        sys.exit(1)

    doc_type = sys.argv[1]
    output_path = sys.argv[2]
    case_data = json.loads(sys.argv[3]) if len(sys.argv) > 3 else {}

    path, cid = generate_doc(doc_type, case_data, output_path)
    print(f'OK|{path}|{cid}')

if __name__ == '__main__':
    main()
